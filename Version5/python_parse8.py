# http://stackoverflow.com/questions/27861732/parsing-of-table-from-docx-file
# 27-07-2017: Test program to parse FedX documents

from docx import Document 
import openpyxl
import sys
#import win32com.client
import os
import re
import logging
import datetime
from datetime import time
from shutil import copyfile 
import itertools #This module is required for iterating through two arrays/objects in parallel. See "copy_across_sheet_ids"

from docx.shared import Pt


def parse_document_table(table):
	# Seems that package info starts at row 22...
	for row_number, row in enumerate(table.rows, 1):		
			
		#print(str(row_number)+'\t'+str(row))
	
		for cell_number, cell in enumerate(table.columns):			
	
			#Try and start from row 41 in document
			#print "row is: %s" % row_number
			#print "cell is: %s" % cell_number
			#print"In the right place?"
		
			cell_value = table.rows[row_number].cells[cell_number].text
		 
			if cell_value != '':
				#strip off excess characters
				#cell_value = cell_value[0:1]
				#print cell_value	
			
				'''
				if cell_value == "No. of Packages":
					#print"How many packages?"
					#print "row number is: %s" % row_number	
				
					#Center quantity field
					quantity = '1'.rjust(8)
				
					#Add quantity to this column
					table.rows[row_number + 1].cells[cell_number].text = quantity	
					
				if cell_value == "No. of Units":
				
				
				
						check_value = table.rows[row_number + 1].cells[cell_number].text
					
						print "Value is: %s" % check_value
				
						#Centre value
						number_of_samples = number_of_samples.rjust(10)				
				
						#For each row, need to convert to decimal 				
				
						#Add number of units to first row
						table.rows[row_number + 1].cells[cell_number].text = number_of_samples
						#Assign total value column value also as number_of_samples
				
						if number_of_samples.find('.') == -1:
							number_of_samples = number_of_samples + '.00'				
				
						number_of_samples = number_of_samples.rjust(45)
						table.rows[row_number + 1].cells[12].text = number_of_samples				
						#Second row
						table.rows[row_number + 4].cells[cell_number].text = number_of_samples				
				
						#Third row
						table.rows[row_number + 7].cells[cell_number].text = number_of_samples		
				
						#Subtotal
						table.rows[32].cells[12].text = number_of_samples
				
						#Invoice Total
						table.rows[43].cells[12].text = number_of_samples
				
				
						#Remove spaces
						number_of_samples = number_of_samples.replace(" ", "")				
				
			
				if "Row" in cell_value:
					print "Found Row3!"
					print "Row at Row3 is: %s" % row_number
					print "Column atRow3 is: %s" % cell_number				
			
			
				if cell_value == "Description of Goods (Including Harmonized Tariff No.)":											
				
					table.rows[row_number + 1].cells[cell_number].text = "%s x %s vials containing < %s ..." % (number_of_samples, vial_size, max_quantity)
				'''
			
				if 'Marker' in cell_value:
					print "Found %s" % cell_value
					print "Row at %s is: %s" % (cell_value, row_number)
					print "Column at %s is: %s" % (cell_value, cell_number)	
					
					#Clear cell value if Marker is in it
					table.rows[row_number].cells[cell_number].text = ''	
				
			
			
				if row_number > 45:
					break			
	
	
		if row_number > 45:
		
			return





#Global Variables 
vial_size = '4 mL'
number_of_samples = '16'
max_quantity = '100 mg'
AWB = '7796 9069 5644'
Date = '20th July 2017'
name = 'Dr Matthew R. Crittall'
freight_value = '65.10'


#Get all FedX docs in folder (should be word_docs)

#Fedx_docs = [each for each in os.listdir("/Users/matthewcrittall/Documents/Programming/Webdev/Python/Projects/Dispatch_Python/dispatch_git/Version5") if each.endswith('.docx')]

#print Fedx_docs

#doc = "S:\Receptos\Team_individual_folders\Matt_C\VBA_test\FedExpy\Commercial_Invoice_Smallpdf.docx"	
doc = "Commercial_Invoice_ilovepdf_txtbox3.docx"			
		
# Try and open word document. See: http://stackoverflow.com/questions/1134607/python-exception-handling and http://www.pythonforbeginners.com/error-handling/python-try-and-except 
try:
	document = Document(doc)
except:		
	# If a file is still open, does not seem to be caught by this exception but rather the one below when saving the document
	print "\nSorry, the file %s is still open. Please close it and try running the script again\n" % doc	
	quit()
	
	
table = document.tables[0]


#Set document styles. See here: https://stackoverflow.com/questions/27884703/set-paragraph-font-in-python-docx
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(7)	


print "There are %s table(s) in this document" % len(document.tables)	

print "The number of rows in table is: %s" % len(table.rows)
print "The number of colunns in table is: %s" % len(table.columns)

#Function to scan and obtain field positions in document if labelled with 'Marker'
parse_document_table(table)


#Date
table.rows[1].cells[7].text = ' ' + Date

#AWB
table.rows[3].cells[7].text = ' ' + AWB

#Name
table.rows[47].cells[10].text = ' ' + name

#Number of packages in row1
quantity = '1'.rjust(8)
table.rows[23].cells[0].text = quantity

#Number of units in row1
number_of_samples = number_of_samples.rjust(10)
table.rows[23].cells[2].text = number_of_samples

#Description of products in row1
#Clear spaces in sample string
number_of_samples = number_of_samples.replace(" ", "")
table.rows[23].cells[5].text = "%s x %s vials containing < %s ..." % (number_of_samples, vial_size, max_quantity)

#Total value in row1
number_of_samples = number_of_samples + '.00'
round(float(number_of_samples), 2)
table.rows[23].cells[12].text = number_of_samples.rjust(43)

#subtotal
table.rows[32].cells[12].text = number_of_samples.rjust(43)

#Add Marker13 (subtotal) and Marker15 (Freight) to get Marker14 (Invoice Total)
invoice_total = float(table.rows[32].cells[12].text) + float(freight_value)
#print invoice_total
#Place value at Marker14
invoice_total = str(invoice_total)
table.rows[43].cells[12].text = invoice_total.rjust(44)





try:						
	document.save("comm_invoice_copy.docx")
except:		
	print "\nSorry, the file %s is still open..." % doc
	

			
		
		
		
	
		
		
		


	
	
	
	
	
		
		