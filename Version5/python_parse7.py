# http://stackoverflow.com/questions/27861732/parsing-of-table-from-docx-file
# 27-07-2017: Test program to parse FedX documents

from docx import Document 
import openpyxl
import sys
#import win32com.client
import os
import re
import logging
import datetime
from datetime import time
from shutil import copyfile 
import itertools #This module is required for iterating through two arrays/objects in parallel. See "copy_across_sheet_ids"

# for adjusting font size
from docx.shared import Pt


# for aligning font
from docx.enum.text import WD_ALIGN_PARAGRAPH


#Variables passed from FedEx object
vial_size = '4 mL'
number_of_samples = '16'
max_quantity = '100 mg'		

#Get all FedX docs in folder (should be word_docs)

Fedx_docs = [each for each in os.listdir("/Users/matthewcrittall/Documents/Programming/Webdev/Python/Projects/Dispatch_Python/dispatch_git/Version5") if each.endswith('.docx')]

print Fedx_docs

doc = "Commercial_Invoice_ilovepdf_txtbox2.docx"		
		
# Try and open word document. See: http://stackoverflow.com/questions/1134607/python-exception-handling and http://www.pythonforbeginners.com/error-handling/python-try-and-except 
try:
	document = Document(doc)
except:		
	# If a file is still open, does not seem to be caught by this exception but rather the one below when saving the document
	print "\nSorry, the file %s is still open. Please close it and try running the script again\n" % doc	
	quit()
	
	
	# Get one and only table
table = document.tables[0]	


#Set document styles. See here: https://stackoverflow.com/questions/27884703/set-paragraph-font-in-python-docx
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(7)	


print "There are %s table(s) in this document" % len(document.tables)	

print "The number of rows in table is: %s" % len(table.rows)
print "The number of colunns in table is: %s" % len(table.columns)

'''
#No.of packages
table.rows[42].cells[1].text = "1"

#No.of units
table.rows[42].cells[8].text = "      16"	

#Description of Goods
table.rows[42].cells[15].text = "Synthetic compounds for use in pharmaceutical studies"

#Description of Goods
table.rows[42].cells[10].text = "Each"

try:						
	document.save("comm_invoice_copy.docx")			

except:		
	print "\nSorry, the file %s is still open..." % doc
	
'''




# Seems that package info starts at row 22...
for row_number, row in enumerate(table.rows, 22):		
			
	#print(str(row_number)+'\t'+str(row))
	
	for cell_number, cell in enumerate(table.columns):			
	
		#Try and start from row 41 in document
		#print "row is: %s" % row_number
		#print "cell is: %s" % cell_number
		#print"In the right place?"
		
		cell_value = table.rows[row_number].cells[cell_number].text
		 
		if cell_value != '':
			#strip off excess characters
			#cell_value = cell_value[0:1]
			#print cell_value
			
			'''
			if "Row2" in cell_value:
				print "Found X!"
				print "Row at X is: %s" % row_number
				print "Column at X is: %s" % cell_number
				#cell.txt = "Testing 101"
				#table.rows[row_number].cells[cell_number].text = ""
				
			if "Row3" in cell_value:
				print "Found Z!"
				print "Row at Z is: %s" % row_number
				print "Column at Z is: %s" % cell_number
				#cell.txt = "Testing 101"
				#table.rows[row_number].cells[cell_number].text = ""
			
			'''
			if cell_value == "No. of Packages":
				#print"How many packages?"
				#print "row number is: %s" % row_number	
				
				#Add quantity to this column
				table.rows[row_number + 1].cells[cell_number].text = "        1"		
				
			if cell_value == "No. of Units":
				#print"How many packages?"
				
				if table.rows[row_number + 1].cells[cell_number].text == '':					
					#Add number of units (with whitespace)
					number_of_samples = '          ' + number_of_samples					
					table.rows[row_number + 1].cells[cell_number].text = number_of_samples
					
					table.rows[row_number + 2].cells[cell_number].text = number_of_samples
					
					#Remove whitespace
					number_of_samples = number_of_samples.replace(" ", "")	
					
						
					
			
			if cell_value == "Description of Goods (Including Harmonized Tariff No.)":			
							
				
				table.rows[row_number + 1].cells[cell_number].text = "%s x %s vials containing < %s ..." % (number_of_samples, vial_size, max_quantity)
			
			
			
			'''
			if cell_value == "C":
				print "Found C!"
				print "Row at C is: %s" % row_number
				print "Column at C is: %s" % cell_number
				break
			'''
	
	if row_number > 45:
		# save document 
		try:						
			document.save("comm_invoice_copy.docx")
			quit()	
	
		except:		
			print "\nSorry, the file %s is still open..." % doc
			quit()
			
		
			
		
		
		
	
		
		
		


	
	
	
	
	
		
		